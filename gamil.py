import imaplib
import email
from email.header import decode_header
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
import pytesseract
from pdf2image import convert_from_path
import spacy
from word2number import w2n
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
import io

nlp = spacy.load("en_core_web_sm")

# Path to Tesseract-OCR executable (adjust for your system)
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\andre\anaconda3\Lib\site-packages\tesseract"

# Sanitize filenames (remove problematic characters)
def sanitize_filename(filename):
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
    sanitized = sanitized.replace('\r', '').replace('\n', '')  # Remove newlines
    sanitized = sanitized.replace('\t', '')  # Remove tabs
    return sanitized

# Extract email body
def extract_email_body(msg):
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            content_disposition = str(part.get("Content-Disposition"))
            if "attachment" not in content_disposition:
                if content_type == "text/plain":
                    return part.get_payload(decode=True).decode("utf-8", errors="ignore")
                elif content_type == "text/html":
                    return part.get_payload(decode=True).decode("utf-8", errors="ignore")
    else:
        return msg.get_payload(decode=True).decode("utf-8", errors="ignore")

# Extract text from DOCX
def extract_text_from_docx(attachment_content):
    doc = Document(io.BytesIO(attachment_content))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Extract text from PDF
def extract_text_from_pdf(attachment_content):
    pdf_reader = PdfReader(io.BytesIO(attachment_content))
    text = ""
    for page in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page].extract_text()
    return text

# OCR for scanned PDFs
def extract_text_with_ocr(pdf_path):
    text = ""
    try:
        images = convert_from_path(pdf_path)
        for image in images:
            text += pytesseract.image_to_string(image, lang='eng')
    except Exception as e:
        print(f"OCR failed for {pdf_path}: {e}")
    return text

# Read resume files
def read_resume_from_file(file_path):
    text = ""
    try:
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            if not text.strip():
                text = extract_text_with_ocr(file_path)
        elif file_path.endswith(".doc"):
            import comtypes.client
            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
        else:
            raise ValueError("Unsupported file format")
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
    return text

# Extract resume details
def extract_name_from_text(text):
    text = text.strip()
    text = re.sub(r'\S+@\S+', '', text)
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    lines = text.split("\n")
    irrelevant_words = ["summary", "contact", "education", "experience", "skills", "references", "profile", "resume", "cv"]
    for line in lines[:3]:
        line = line.strip()
        if any(irrelevant_word in line.lower() for irrelevant_word in irrelevant_words):
            continue
        if len(line) > 1:
            name_parts = line.split()
            if len(name_parts) > 1:
                return " ".join([part.title() for part in name_parts])
            elif len(name_parts) == 1:
                return name_parts[0].title()
    return "Name not found"

# Function to extract email from resume text
def extract_email_from_text(text):
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return email_match.group(0) if email_match else "Email not found"

# Function to extract phone numbers
def extract_phone_from_text(text):
    phone_pattern = re.compile(r"(?:direct|mobile|phone|ph#|contact|tel|cell)?[:\s-]*"
                               r"(?:\+?\d{1,3}[-.\s]?)?"
                               r"\(?\d{1,4}\)?"
                               r"[-.\s]?\d{1,4}"
                               r"[-.\s]?\d{1,4}"
                               r"[-.\s]?\d{1,9}"
                               r"(?:\s?(?:ext|x|extension)\s?\d{1,5})?")
    matches = phone_pattern.findall(text)
    phones = [re.sub(r"[^+\d\s()-]", "", match).strip() for match in matches if len(re.sub(r"\D", "", match)) >= 10]
    return ", ".join(phones) if phones else "Phone not found"

# Function to extract experience from resume text
def extract_experience(text):
    text = text.lower()
    numeric_pattern = r"(?:more than|over|at least|around|approximately|nearly|up to)?\s*(\d+)\+?\s*years?"
    numeric_match = re.search(numeric_pattern, text)
    if numeric_match:
        years = numeric_match.group(1)
        return f"{int(years)}+ years" if '+' in numeric_match.group(0) else f"{int(years)} years"
    return "Experience not found"

# Function to extract all skills from the resume mentioned in the job description
def extract_relevant_skills(resume_text, job_desc_subject):
    """
    Extract skills from job description subject and check if they exist in the resume.
    The subject will only mention skills after specific keywords like 'with'.
    """
    # Extract skills after 'with'
    skill_pattern = re.compile(r'with (.+)', re.IGNORECASE)
    match = skill_pattern.search(job_desc_subject)
    
    if not match:
        return []
    
    # Get skills list from subject
    skills_str = match.group(1).strip()
    subject_skills = [skill.strip().lower() for skill in re.split(r',|/', skills_str)]
    
    # Prepare resume text for matching
    resume_lower = resume_text.lower()
    
    # Find matches with whole word matching
    matched_skills = []
    for skill in subject_skills:
        pattern = re.compile(rf'\b{re.escape(skill)}\b', re.IGNORECASE)
        if pattern.search(resume_lower):
            matched_skills.append(skill.title())
    
    return matched_skills

# Function to extract certifications
def extract_certifications_count(text):
    certification_keywords = [
        r"certification", r"certifications", r"certified", r"certificate", r"certificates"
    ]
    pattern = r"|".join(certification_keywords)
    matches = re.findall(pattern, text, re.IGNORECASE)
    return len(matches)

# Function to extract location from resume text
def extract_location_from_text(text):
    """Extract location (city, state, or ZIP code) from resume text."""
    location_match = re.search(
        r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s(?:TX|CA|NY|FL|WA|IL|PA|GA|NC|OH|NJ|VA|CO|AZ|MA|MD|TN|MO|IN|WI|MN|SC|AL|LA|KY|OR|OK|CT|IA|MS|KS|AR|NV|UT|NM|NE|WV|ID|HI|ME|NH|MT|RI|DE|SD|ND|AK|VT|WY))\b"  # City, State
        r"|\b\d{5}(?:-\d{4})?\b",  # ZIP code
        text
    )
    if location_match:
        location = location_match.group(0)
        if not any(keyword in location.lower() for keyword in ["assistant", "server", "sql"]):  # Example of filtering out unrelated matches
            return location
    return "Location not found"

# Extract government from resume text
def extract_government_details(text):
    """
    Extract the first current working location from the text using multiple patterns for 'Present', 'Till Date', etc.
    Removes unwanted prefixes and extra whitespace.
    """
    # Patterns to detect blocks with work location details
    patterns = [
        r"(Client:.*?Present|Client:.*?\d{4}|Client:.*?Till Date)",  # Client and its timeframe
        r"(Professional Experience:.*?Present|Professional Experience:.*?\d{4}|Professional Experience:.*?Till Date)",
        r"(EXPERIENCE.*?Present|EXPERIENCE.*?\d{4}|EXPERIENCE.*?Till Date)",
        r"(Past work:.*?Present|Past work:.*?\d{4}|Past work:.*?Till Date)",
        r"(WORK EXPERIENCE:.*?Present|WORK EXPERIENCE:.*?\d{4}|WORK EXPERIENCE:.*?Till Date)",
    ]
   
    # Extract the relevant sections based on the patterns
    extracted_sections = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL)
        extracted_sections.extend(matches)
    
    # Combine all extracted sections into one string for further processing
    combined_text = " ".join(extracted_sections)
    
    # Define the combined pattern to extract location information before keywords like 'Present', 'Till Date', etc.
    location_pattern = re.compile(
        r"""
        # Flexible location matching with optional "Client:" prefix and keywords like 'Present', 'Till Date'
        (?:Client:\s*)?                                      # Optional 'Client:' prefix
        ([A-Za-z\s,.()]+(?:USA|México|Virginia|FL|NJ|Texas|Tallahassee|Reston|New York|U\.S\.A\.|U\.S\.|America))  # Location
        .*?                                                  # Any text in between
        (?=\s*(?:Present|Till Date|to date|current|\d{4}[-–]\d{4}|[\w\s]+))  # Lookahead for keywords or date patterns
        
        |  # OR
        
        # Stricter format where "Client:" is explicitly present and followed by "Present"
        Client:\s*                                            # 'Client:' prefix
        ([A-Za-z\s,]+)                                        # Location
        \s+[A-Z][a-z]+\s\d{4}\s*[-—]\s*Present                # Date range ending with 'Present'
        """,
        re.IGNORECASE | re.VERBOSE
    )
   
    # Find the first match for locations within the extracted sections
    match = location_pattern.search(combined_text)
    
    # Check if a match is found
    if match:
        # Extract the first location
        first_location = match.group(0).strip()
        
        # Remove unwanted parts like "Client:", extra whitespace, and any date/time information
        cleaned_location = re.sub(r"(Client:|Present|EXPERIENCE|Past work:|WORK EXPERIENCE:|\d{4}[-–]\d{4}|[A-Za-z]+\s\d{4}\s*[-—]\s*Present|[\t\n]+)", "", first_location)
        cleaned_location = re.sub(r"\s{2,}", " ", cleaned_location).strip()  # Remove extra spaces
        
        # Format the result
        formatted_location = f"[{cleaned_location}]"
        return formatted_location
    else:
        # If no matches, return a default "Not found"
        return "Not found"

# Function to extract visa status from the resume text
def extract_visa_status(text):
    """Extract visa status from the resume text."""
    visa_keywords = {
        "H1B": ["h1b"],
        "Green Card": ["green card", "permanent resident"],
        "US Citizen": ["usc", "us citizen", "citizenship: us"],
        "OPT": ["opt"],
        "CPT": ["cpt"],
        "L2": ["l2 visa"],
        "EAD": ["ead"],
        "TN Visa": ["tn visa"],
        "Study Visa": ["study visa"]
    }
    visa_status = []
    for visa, patterns in visa_keywords.items():
        for pattern in patterns:
            if re.search(pattern, text.lower()):
                visa_status.append(visa)
                break
    return ", ".join(visa_status) if visa_status else "Not found"

# Calculating resume score
def calculate_resume_score(resume_text, job_desc_text, skills, experience, certifications, visa_status, location, government):
    corpus = [job_desc_text, resume_text]
    vectorizer = CountVectorizer().fit_transform(corpus)
    vectors = vectorizer.toarray()

    # Cosine Similarity: Measures how closely the resume text aligns with the job description.
    similarity_score = cosine_similarity([vectors[0]], [vectors[1]])[0][0]

    # Skills, experience, and certifications
    skills_count = len(skills)
    experience_years = int(re.search(r"\d+", experience).group(0)) if re.search(r"\d+", experience) else 0
    certifications_count = certifications
 
    normalized_experience = min(experience_years / 20, 1)
    normalized_skills = min(skills_count / 20, 1)
 
    # Visa Status Scoring
    visa_priority = {
        "US Citizen": 1.0,
        "Green Card": 0.9,
        "H1B": 0.8,
        "OPT": 0.7,
        "CPT": 0.6,
        "L2": 0.5,
        "EAD": 0.5,
        "TN Visa": 0.6,
        "Not found": 0.0
    }
    visa_score = visa_priority.get(visa_status, 0.0)
 
    # Location Scoring
    location_score = 0.0
    if location.lower() != "location not found":
        location_score = 1.0  # Location provided gets full credit
 
    # Government Scoring
    government_score = 0.0
    if government.lower() != "not found":
        government_score = 1.0  # Government-related experience gets full credit
 
    # Weighted scoring
    score = (
        similarity_score * 0.5 +  # Adjusted to 40% weight
        normalized_skills * 0.8 +  # Adjusted to 25% weight
        normalized_experience * 0.01 +  # Adjusted to 25% weight
        certifications_count * 0.01 +  # Certifications contribute 20%
        visa_score * 0.05 +  # Visa status contributes 5%
        location_score * 0.05 +  # Location contributes 5%
        government_score * 0.05  # Government experience contributes 10%
    )
 
    return round(min(score * 100, 100), 2)

# Function to filter emails by Job ID in subject or body
def filter_emails_by_job_id(job_id, email_ids, mail):
    filtered_emails = []
    for email_id in email_ids:
        status, msg_data = mail.fetch(email_id, "(RFC822)")
        for response_part in msg_data:
            if isinstance(response_part, tuple):
                msg = email.message_from_bytes(response_part[1])
                subject, encoding = decode_header(msg["Subject"])[0]
                if isinstance(subject, bytes):
                    subject = subject.decode(encoding or "utf-8")
                # Match Job ID in Subject and Body (case insensitive)
                if job_id.lower() in subject.lower():
                    filtered_emails.append(msg)
                else:
                    body = extract_email_body(msg)
                    if body and job_id.lower() in body.lower():
                        filtered_emails.append(msg)
    return filtered_emails

# Process resumes and filter emails based on Job ID
def process_resumes_and_attachments(job_id):
    mail = imaplib.IMAP4_SSL('imap.gmail.com')
    mail.login('k12392945@gmail.com', 'xcya gowp wxrd cjav')  # Secure your password
    mail.select("inbox")

    print(f"Processing emails for Job ID: {job_id}")
    status, messages = mail.search(None, 'ALL')
    email_ids = messages[0].split()

    # Filter emails by Job ID
    filtered_emails = filter_emails_by_job_id(job_id, email_ids, mail)
    print(f"Found {len(filtered_emails)} emails matching the Job ID: {job_id}")

    resume_details = []

    for msg in filtered_emails:
        job_desc_text = extract_email_body(msg)

        for part in msg.walk():
            if part.get("Content-Disposition") and "attachment" in part.get("Content-Disposition"):
                attachment_filename = sanitize_filename(part.get_filename())
                attachment_content = part.get_payload(decode=True)

                if attachment_filename.lower().endswith('.pdf'):
                    resume_text = extract_text_from_pdf(attachment_content)
                elif attachment_filename.lower().endswith('.docx'):
                    resume_text = extract_text_from_docx(attachment_content)
                else:
                    continue

                # Extract details from the resume text
                details = {
                    "name": extract_name_from_text(resume_text),
                    "email": extract_email_from_text(resume_text),
                    "phone": extract_phone_from_text(resume_text),
                    "experience": extract_experience(resume_text),
                    "skills": extract_relevant_skills(resume_text, job_desc_text),
                    "certifications": extract_certifications_count(resume_text),
                    "location": extract_location_from_text(resume_text),  
                    "visa_status": extract_visa_status(resume_text),  
                    "government": extract_government_details(resume_text)
                }

                # Calculate the resume score
                score = calculate_resume_score(resume_text, job_desc_text, details['skills'],
                                               details['experience'], details['certifications'],
                                               details['visa_status'], details['location'], details['government'])
                details['Resume Score'] = score
                resume_details.append(details)

    mail.logout()

    # Create DataFrame from the extracted resume details
    df = pd.DataFrame(resume_details)

    # Create rank function using resume score
    def assign_rank(score):
        if 0 <= score <= 9:
            return 10
        elif 10 <= score <= 19:
            return 9
        elif 20 <= score <= 29:
            return 8
        elif 30 <= score <= 39:
            return 7
        elif 40 <= score <= 49:
            return 6
        elif 50 <= score <= 59:
            return 5
        elif 60 <= score <= 69:
            return 4
        elif 70 <= score <= 79:
            return 3
        elif 80 <= score <= 89:
            return 2
        elif 90 <= score <= 100:
            return 1
        return 10  # Default if something goes wrong

        # Apply rank assignment based on the score
    df['Rank'] = df['Resume Score'].apply(assign_rank)  # Assign ranks
    df = df.sort_values(by="Rank", ascending=True)  # Sort in ascending order

    return df

# Example usage:
job_id = input("Enter the Job ID to search: ")
df = process_resumes_and_attachments(job_id)
df